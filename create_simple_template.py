from docx import Document
from pathlib import Path

# Create a simple template WITHOUT dynamic tables for testing
template_path = Path("examples/Project_Sample/Template/template_simple.docx")
template_path.parent.mkdir(parents=True, exist_ok=True)

doc = Document()
doc.add_heading('HOP DONG MAU', 0)
doc.add_paragraph('Cong Hoa Xa Hoi Chu Nghia Viet Nam')
doc.add_paragraph('Doc lap - Tu do - Hanh phuc')
doc.add_paragraph('---')
doc.add_paragraph('So Hop Dong: <<SoHD>>')
doc.add_paragraph('Ben A: Cong Ty ABC')
doc.add_paragraph('Ben B: <<HoTen>>')
doc.add_paragraph('Dia Chi: <<DiaChi>>')

doc.add_heading('Thong Tin Them', level=1)
doc.add_paragraph('ID: <<ID>>')

doc.save(template_path)
print(f"Created simple template: {template_path}")
print("This template only uses simple placeholders, no dynamic tables.")
print("Update config.yaml to use 'Template/template_simple.docx' to test.")
