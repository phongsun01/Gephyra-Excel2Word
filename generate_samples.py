import pandas as pd
from docx import Document
from pathlib import Path

# Paths
base_dir = Path("examples/Project_Sample")
data_dir = base_dir / "Data"
template_dir = base_dir / "Template"
output_dir = base_dir / "Output"

data_dir.mkdir(parents=True, exist_ok=True)
template_dir.mkdir(parents=True, exist_ok=True)
output_dir.mkdir(parents=True, exist_ok=True)

# 1. Create Sample Excel
excel_path = data_dir / "sample_data.xlsx"

# Main Text Data
df_text = pd.DataFrame([
    {"ID": "001", "HoTen": "Nguyen Van A", "SoHD": "HD001", "DiaChi": "Ha Noi"},
    {"ID": "002", "HoTen": "Tran Thi B", "SoHD": "HD002", "DiaChi": "Ho Chi Minh"},
    {"ID": "003", "HoTen": "Le Van C", "SoHD": "HD003", "DiaChi": "Da Nang"},
])

# Config Data
df_config = pd.DataFrame([
    {"TableName": "BangGia", "SourceSheet": "PriceList", "Range": "A1:C4", "WordBookmark": "TablePos"},
])

# Table Data (PriceList)
df_price = pd.DataFrame([
    {"TenSP": "San Pham 1", "DVT": "Cai", "Gia": 100000},
    {"TenSP": "San Pham 2", "DVT": "Hop", "Gia": 200000},
    {"TenSP": "San Pham 3", "DVT": "Bo", "Gia": 150000},
])

# Write to Excel
with pd.ExcelWriter(excel_path) as writer:
    df_text.to_excel(writer, sheet_name="Text", index=False)
    df_config.to_excel(writer, sheet_name="Config", index=False)
    df_price.to_excel(writer, sheet_name="PriceList", index=False)

print(f"Created Excel: {excel_path}")

# 2. Create Sample Template
template_path = template_dir / "template_hd.docx"
doc = Document()
doc.add_heading('HOP DONG MAU', 0)
doc.add_paragraph('Cong Hoa Xa Hoi Chu Nghia Viet Nam')
doc.add_paragraph('Doc lap - Tu do - Hanh phuc')
doc.add_paragraph('---')
doc.add_paragraph('So Hop Dong: <<SoHD>>')
doc.add_paragraph('Ben A: Cong Ty ABC')
doc.add_paragraph('Ben B: <<HoTen>>')
doc.add_paragraph('Dia Chi: <<DiaChi>>')

doc.add_heading('Danh Sach San Pham', level=1)
doc.add_paragraph('Bang gia chi tiet:')

# Create table with header row
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Ten SP'
hdr_cells[1].text = 'DVT'
hdr_cells[2].text = 'Gia'

# Add template row with docxtpl loop syntax
# docxtpl requires the loop to wrap the row in XML, but python-docx can't do that easily
# So we add a row and manually edit text to include loop tags
row_cells = table.add_row().cells
row_cells[0].text = '<%tr for item in BangGia%><<item.TenSP>>'
row_cells[1].text = '<<item.DVT>>'
row_cells[2].text = '<<item.Gia>><%endtr%>'

doc.save(template_path)
print(f"Created Template: {template_path}")
